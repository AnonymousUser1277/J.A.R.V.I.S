"""
Document Generation System
Creates professional documents from natural language
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict
from datetime import datetime

logger = logging.getLogger(__name__)

# Document templates
try:
    from docx import Document as WordDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("⚠️ Document generation not available. Install: pip install python-docx")


class DocumentGenerator:
    """
    Generate professional documents
    
    Supported formats:
    - Word (.docx)
    - Markdown (.md)
    - Plain text (.txt)
    """
    
    def __init__(self):
        self.output_dir = Path.home() / "Documents" / "JARVIS_Documents"
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_report(
        self,
        title: str,
        content: str,
        format: str = 'docx',
        metadata: Optional[Dict] = None
    ) -> Optional[str]:
        """
        Generate a document report
        
        Args:
            title: Document title
            content: Document content (can be AI-generated)
            format: Output format (docx, md, txt)
            metadata: Additional metadata (author, date, etc.)
            
        Returns:
            Path to generated document
        """
        if format == 'docx':
            return self._generate_docx(title, content, metadata)
        elif format == 'md':
            return self._generate_markdown(title, content, metadata)
        elif format == 'txt':
            return self._generate_text(title, content, metadata)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _generate_docx(self, title: str, content: str, metadata: Optional[Dict]) -> Optional[str]:
        """Generate Word document"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return None
        
        try:
            doc = WordDocument()
            
            # Add title
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            if metadata:
                author = metadata.get('author', 'JARVIS')
                date = metadata.get('date', datetime.now().strftime('%B %d, %Y'))
                
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta_run = meta_para.add_run(f"By {author} | {date}")
                meta_run.font.size = Pt(10)
                meta_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Spacing
            
            # Add content
            # Split by sections (headings)
            sections = self._parse_sections(content)
            
            for section_title, section_content in sections:
                if section_title:
                    doc.add_heading(section_title, 1)
                
                # Add paragraphs
                for paragraph in section_content.split('\n\n'):
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{title}_{timestamp}.docx"
            filepath = self.output_dir / filename
            
            doc.save(str(filepath))
            
            logger.info(f"✅ Document created: {filename}")
            
            # Open the document
            os.startfile(str(filepath))
            
            return str(filepath)
        
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            return None
    
    def _generate_markdown(self, title: str, content: str, metadata: Optional[Dict]) -> str:
        """Generate Markdown document"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{title}_{timestamp}.md"
        filepath = self.output_dir / filename
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                # Title
                f.write(f"# {title}\n\n")
                
                # Metadata
                if metadata:
                    author = metadata.get('author', 'JARVIS')
                    date = metadata.get('date', datetime.now().strftime('%B %d, %Y'))
                    f.write(f"**Author:** {author}  \n")
                    f.write(f"**Date:** {date}  \n\n")
                
                f.write("---\n\n")
                
                # Content
                f.write(content)
            
            logger.info(f"✅ Markdown created: {filename}")
            os.startfile(str(filepath))
            
            return str(filepath)
        
        except Exception as e:
            logger.error(f"Failed to generate Markdown: {e}")
            return None
    
    def _generate_text(self, title: str, content: str, metadata: Optional[Dict]) -> str:
        """Generate plain text document"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{title}_{timestamp}.txt"
        filepath = self.output_dir / filename
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"{title}\n")
                f.write("=" * len(title) + "\n\n")
                
                if metadata:
                    author = metadata.get('author', 'JARVIS')
                    date = metadata.get('date', datetime.now().strftime('%B %d, %Y'))
                    f.write(f"Author: {author}\n")
                    f.write(f"Date: {date}\n\n")
                
                f.write(content)
            
            logger.info(f"✅ Text document created: {filename}")
            os.startfile(str(filepath))
            
            return str(filepath)
        
        except Exception as e:
            logger.error(f"Failed to generate text: {e}")
            return None
    
    def _parse_sections(self, content: str) -> list:
        """Parse content into sections"""
        sections = []
        current_section = (None, "")
        
        for line in content.split('\n'):
            # Check if line is a heading (starts with #)
            if line.startswith('#'):
                if current_section[1]:
                    sections.append(current_section)
                
                # Extract heading
                heading = line.lstrip('#').strip()
                current_section = (heading, "")
            else:
                current_section = (current_section[0], current_section[1] + line + '\n')
        
        if current_section[1]:
            sections.append(current_section)
        
        return sections


# Global instance
_document_generator = None

def get_document_generator() -> DocumentGenerator:
    """Get or create global document generator"""
    global _document_generator
    if _document_generator is None:
        _document_generator = DocumentGenerator()
    return _document_generator


def generate_document_from_prompt(prompt: str, ai_client) -> Optional[str]:
    """
    Generate document from natural language prompt
    
    Example: "Create a report about quarterly sales performance"
    """
    from ai.providers import call_ai_model
    
    # Extract document type and topic
    doc_type = "report"  # Default
    if "memo" in prompt.lower():
        doc_type = "memo"
    elif "letter" in prompt.lower():
        doc_type = "letter"
    elif "proposal" in prompt.lower():
        doc_type = "proposal"
    
    # Generate content using AI
    ai_prompt = f"""
Generate a professional {doc_type} with the following requirements:

{prompt}

Format the output in Markdown with clear headings and sections.
Include:
- An introduction
- Main body with detailed sections
- A conclusion
- Professional tone

Use proper Markdown formatting (# for headings, ** for bold, etc.)
"""
    
    try:
        content = call_ai_model(ai_prompt, ai_client)
        
        if not content:
            return None
        
        # Extract title from first heading
        lines = content.split('\n')
        title = "Untitled Document"
        for line in lines:
            if line.startswith('#'):
                title = line.lstrip('#').strip()
                break
        
        # Generate document
        generator = get_document_generator()
        filepath = generator.generate_report(
            title=title,
            content=content,
            format='docx',
            metadata={'author': 'JARVIS Assistant'}
        )
        
        return filepath
    
    except Exception as e:
        logger.error(f"Document generation failed: {e}")
        return None